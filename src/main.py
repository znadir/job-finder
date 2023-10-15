import tkinter as tk
from tkinter import messagebox, PhotoImage
import os
import docx
from utils import add_hyperlink
from searchJobs import searchJobs
from selenium import webdriver
from selenium.webdriver.firefox.options import Options

class App:
    def __init__(self):
        self.doc_path = os.path.join(os.path.dirname(__file__), "./jobs.docx")

        self.init_gui()

    def init_gui(self):
        """
        Initialize the GUI
        """
        self.root = tk.Tk()
        self.root.geometry("550x230") 
        self.root.title("Job Finder") 
        self.root.resizable(False, False)

        self.app_title = tk.Label(self.root, text="Job Finder", font=("Arial Bold", 17))
        self.app_title.pack(pady=(10, 0))

        self.app_desc = tk.Label(self.root, text="Générer un fichier word avec des jobs pour vous", font=("Arial", 12))
        self.app_desc.pack()

        self.frame = tk.Frame(self.root, pady=10)
        self.frame.columnconfigure(0, weight=1, pad=10)
        self.frame.columnconfigure(1, weight=1, pad=10)
        self.frame.columnconfigure(2, weight=1, pad=10)
        self.frame.columnconfigure(3, weight=1, pad=10)

        self.label_mot_cles = tk.Label(self.frame, text="Rechercher un emploi", font=("Arial", 12))
        self.label_mot_cles.grid(row=0, column=0, sticky="w")

        self.mots_cles_var = tk.StringVar(self.frame)
        self.mots_cles = tk.Entry(self.frame, width=25, font=("Arial", 12), textvariable=self.mots_cles_var)
        self.mots_cles.insert(0, "Stage programmation informatique")
        self.mots_cles.grid(row=0, column=1)

        self.label_location = tk.Label(self.frame, text="Location", font=("Arial", 12))
        self.label_location.grid(row=1, column=0, sticky="w")

        self.location_var = tk.StringVar(self.frame)
        self.location = tk.Entry(self.frame, width=25, font=("Arial", 12), textvariable=self.location_var)
        self.location.insert(0, "Laval")
        self.location.grid(row=1, column=1)

        self.label_must_mots = tk.Label(self.frame, text="Doit inclure", font=("Arial", 12))
        self.label_must_mots.grid(row=2, column=0, sticky="w")

        self.must_mots_var = tk.StringVar(self.frame)
        self.must_mots = tk.Entry(self.frame, width=25, font=("Arial", 12), textvariable=self.must_mots_var)
        self.must_mots.insert(0, "Stage")
        self.must_mots.grid(row=2, column=1)

        self.label_headless = tk.Label(self.frame, text="Is headless", font=("Arial", 12))
        self.label_headless.grid(row=3, column=0, sticky="w")

        self.is_headless_var = tk.BooleanVar(self.frame)
        self.is_headless = tk.Checkbutton(self.frame, variable=self.is_headless_var)
        self.is_headless.grid(row=3, column=1, sticky="w")

        self.frame.pack()

        self.frame_btn = tk.Frame(self.root, pady=10)
        self.frame_btn.columnconfigure(0, weight=1)

        self.generate_file = tk.Button(self.frame_btn, text="Generate File", width=10, background="#183cde", fg="#ffffff", command=self.generate_file)
        self.generate_file.grid(row=3, column=0)

        self.open_file = tk.Button(self.frame_btn, text="Open File", width=10, background="#141414", fg="#ffffff", command=self.open_file)
        self.open_file.grid(row=3, column=1)

        self.frame_btn.pack()

        self.root.mainloop()
    
    def create_document(self, query: str, my_jobs: list):
        """
        Create a document with the jobs found
        """
        doc = docx.Document()
        doc.add_heading(f"Recherche d'emploi: {query}", 0)

        for job in my_jobs:
            p = doc.add_paragraph()
            p.add_run(job['title']).bold = True
            doc.add_paragraph(job['snippet'])
            p = doc.add_paragraph()
            add_hyperlink(p, "Voir l'offre emploi", job['link'])
            doc.add_paragraph()

        doc.save(self.doc_path)

    def generate_file(self):
        location = self.location_var.get()
        query = self.mots_cles_var.get()
        must_keywords = self.must_mots_var.get()
        must_keywords = must_keywords.replace(" ", "").split(",")

        options = Options()

        if self.is_headless_var.get():
            options.add_argument("--headless")

        driver = webdriver.Firefox(options=options)

        my_jobs = []

        my_jobs += searchJobs(driver, "https://emplois.ca.indeed.com/jobs?q=", '.jobsearch-ResultsList > li', '.jobTitle span', '.jobTitle > a', '.job-snippet', query, must_keywords)
        my_jobs += searchJobs(driver, f"https://ca.jooble.org/SearchResult?rgns={location}%2C%20QC&ukw=", "div > article", "a", "a", "section > div > div", query, must_keywords)
        my_jobs += searchJobs(driver, f"https://www.option-carriere.ca/recherche/emplois?l={location}&radius=10&sort=relevance&s=", "ul.jobs li", "header h2 a", "header h2 a", "div.desc", query, must_keywords)

        self.create_document(query, my_jobs)

        driver.close()
        messagebox.showinfo("Job Finder", "File generated successfully")
    
    def open_file(self):
        if os.path.exists(self.doc_path):
            os.startfile(self.doc_path)
        else:
            messagebox.showerror("Job Finder", "File not found. Please generate it first")

App()